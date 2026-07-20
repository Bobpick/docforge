"""DOCX (Word) document handler — paragraphs, tables, images, styles."""

import re
import io
import zipfile
from pathlib import Path
from typing import Tuple, Dict, Any, List, Optional

from .utils import (
    format_table_md,
    format_table_json,
    detect_math_expression,
    clean_whitespace,
    generate_image_filename,
    is_duplicate_image,
    compute_hash,
)


class DOCXHandler:
    """Convert DOCX files to Markdown and structured JSON."""

    def __init__(self, extract_images: bool = True, extract_tables: bool = False):
        self.extract_images = extract_images
        self.extract_tables = extract_tables
        self._seen_image_hashes: set = set()

    def convert(
        self, file_path: Path
    ) -> Tuple[str, Dict[str, Any], List[Dict[str, Any]], Dict[str, Any]]:
        """Convert a DOCX file."""
        try:
            from docx import Document
            from docx.oxml.ns import qn
        except ImportError:
            raise ImportError(
                "python-docx is required for DOCX conversion. "
                "Install it with: pip install python-docx"
            )

        doc = Document(str(file_path))
        metadata = self._extract_metadata(doc)

        # Extract images first so we can reference them
        images = []
        if self.extract_images:
            images = self._extract_images(doc, file_path)

        # Walk through document body elements in order
        sections: List[Dict[str, Any]] = []
        md_parts: List[str] = []

        for element in doc.element.body:
            tag = element.tag.split("}")[-1] if "}" in element.tag else element.tag

            if tag == "p":
                # Paragraph
                para = None
                for p in doc.paragraphs:
                    if p._element is element:
                        para = p
                        break
                if para is None:
                    continue

                md_text, sec = self._process_paragraph(para, len(images))
                if md_text.strip():
                    md_parts.append(md_text)
                    if sec:
                        sections.append(sec)

            elif tag == "tbl":
                if not self.extract_tables:
                    continue
                table = None
                for t in doc.tables:
                    if t._element is element:
                        table = t
                        break
                if table is None:
                    continue

                md_text, sec = self._process_table(table)
                if md_text.strip():
                    md_parts.append(md_text)
                    sections.append(sec)

        markdown = clean_whitespace("\n\n".join(md_parts))
        structured = self._build_structured(sections, metadata, images)
        return markdown, structured, images, metadata

    # ------------------------------------------------------------------
    # Metadata
    # ------------------------------------------------------------------
    def _extract_metadata(self, doc) -> Dict[str, Any]:
        core = doc.core_properties
        return {
            "title": core.title or "",
            "author": core.author or "",
            "subject": core.subject or "",
            "created": str(core.created) if core.created else "",
            "modified": str(core.modified) if core.modified else "",
            "page_count": len(doc.sections),
        }

    # ------------------------------------------------------------------
    # Paragraph processing
    # ------------------------------------------------------------------
    def _process_paragraph(
        self, para, image_count: int
    ) -> Tuple[str, Optional[Dict[str, Any]]]:
        """Process a single paragraph element."""
        from docx.oxml.ns import qn

        style_name = para.style.name if para.style else ""
        style_lower = style_name.lower()

        # --- Check for inline images ---
        has_image = False
        for run in para.runs:
            drawing_elements = run._element.findall(qn("w:drawing"))
            inline_shapes = run._element.findall(qn("w:pict"))
            if drawing_elements or inline_shapes:
                has_image = True
                break

        # --- Empty paragraph ---
        text = para.text.strip()
        if not text and not has_image:
            return "", None

        # --- Headings ---
        if style_lower.startswith("heading"):
            try:
                level = int(style_lower.replace("heading", "").strip())
            except ValueError:
                level = 2
            level = max(1, min(6, level))
            md_text = f"{'#' * level} {text}"
            return md_text, {"type": "heading", "level": level, "content": text}

        # --- Title ---
        if "title" in style_lower:
            md_text = f"# {text}"
            return md_text, {"type": "heading", "level": 1, "content": text}

        # --- List items ---
        if style_lower.startswith("list"):
            prefix = self._get_list_prefix(para)
            # Format runs with inline formatting
            formatted = self._format_runs(para.runs)
            md_text = f"{prefix} {formatted}"
            return md_text, {"type": "list_item", "content": formatted, "prefix": prefix}

        # --- Code blocks (by style name or monospace font) ---
        is_code_style = any(
            kw in style_lower
            for kw in ["code", "source", "preformatted", "no spacing", "html preformatted"]
        )
        is_mono_para = self._is_mono_paragraph(para)
        if is_code_style or is_mono_para:
            md_text = f"```\n{text}\n```"
            return md_text, {"type": "code", "content": text}

        # --- Blockquote ---
        if "quote" in style_lower:
            md_text = "> " + text.replace("\n", "\n> ")
            return md_text, {"type": "quote", "content": text}

        # --- Math detection ---
        if detect_math_expression(text):
            if "\n" in text:
                md_text = f"$$\n{text}\n$$"
            else:
                md_text = f"${text}$"
            return md_text, {"type": "math", "content": text}

        # --- Regular paragraph with formatting ---
        formatted = self._format_runs(para.runs)
        if has_image:
            img_ref = f"![Image](images/image_{image_count + 1:03d}.png)"
            formatted = f"{formatted}\n\n{img_ref}" if formatted else img_ref

        return formatted, {"type": "paragraph", "content": formatted}

    def _format_runs(self, runs) -> str:
        """Format paragraph runs with inline markdown (bold, italic, code)."""
        parts = []
        for run in runs:
            text = run.text
            if not text:
                continue

            is_bold = run.bold
            is_italic = run.italic
            font_name = run.font.name or ""
            is_mono = any(
                m in font_name.lower()
                for m in ["consol", "courier", "mono", "code"]
            )

            if is_mono:
                parts.append(f"`{text}`")
            elif is_bold and is_italic:
                parts.append(f"***{text}***")
            elif is_bold:
                parts.append(f"**{text}**")
            elif is_italic:
                parts.append(f"*{text}*")
            else:
                parts.append(text)

        return "".join(parts)

    def _is_mono_paragraph(self, para) -> bool:
        """Check if a paragraph is predominantly monospace font (code block)."""
        if not para.runs:
            return False
        mono_count = 0
        total_runs = 0
        for run in para.runs:
            if not run.text.strip():
                continue
            total_runs += 1
            font_name = (run.font.name or "").lower()
            if any(m in font_name for m in ["consol", "courier", "mono", "code"]):
                mono_count += 1
        # If most runs use monospace, treat as code block
        return total_runs > 0 and mono_count / total_runs >= 0.8

    def _get_list_prefix(self, para) -> str:
        """Determine the list prefix (bullet or numbered)."""
        style_lower = para.style.name.lower() if para.style else ""
        if "number" in style_lower:
            # Try to extract number from text
            match = re.match(r"(\d+)[.)]", para.text.strip())
            if match:
                return f"{match.group(1)}."
            return "1."
        return "-"

    # ------------------------------------------------------------------
    # Table processing
    # ------------------------------------------------------------------
    def _process_table(self, table) -> Tuple[str, Dict[str, Any]]:
        """Process a Word table into Markdown and JSON."""
        rows_data = []
        for row in table.rows:
            cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
            rows_data.append(cells)

        if not rows_data:
            return "", {"type": "table", "headers": [], "rows": []}

        # First row as headers
        headers = rows_data[0]
        rows = rows_data[1:] if len(rows_data) > 1 else []

        md_text = format_table_md(headers, rows)
        return md_text, format_table_json(headers, rows)

    # ------------------------------------------------------------------
    # Image extraction
    # ------------------------------------------------------------------
    def _extract_images(self, doc, file_path: Path) -> List[Dict[str, Any]]:
        """Extract embedded images from a DOCX file (which is a ZIP archive)."""
        images = []

        try:
            with zipfile.ZipFile(str(file_path), "r") as zf:
                media_files = [
                    f for f in zf.namelist()
                    if f.startswith("word/media/") and not f.endswith("/")
                ]

                for idx, media_file in enumerate(media_files, 1):
                    img_data = zf.read(media_file)

                    # Skip tiny images and duplicates
                    if len(img_data) < 500:
                        continue
                    if is_duplicate_image(img_data, self._seen_image_hashes):
                        continue

                    # Determine extension
                    ext = Path(media_file).suffix.lstrip(".") or "png"
                    if ext == "emf" or ext == "wmf":
                        ext = "png"  # Can't easily convert these; save as placeholder

                    filename = f"image_{idx:03d}.{ext}"
                    images.append(
                        {
                            "filename": filename,
                            "data": img_data,
                            "extension": ext,
                            "page": 0,
                            "source": media_file,
                        }
                    )

        except Exception:
            pass

        return images

    # ------------------------------------------------------------------
    # Structured JSON
    # ------------------------------------------------------------------
    def _build_structured(
        self,
        sections: List[Dict[str, Any]],
        metadata: Dict[str, Any],
        images: List[Dict[str, Any]],
    ) -> Dict[str, Any]:
        """Build the structured JSON representation."""
        # Add image references to structured data
        for img in images:
            sections.append(
                {
                    "type": "image",
                    "path": f"images/{img['filename']}",
                    "filename": img["filename"],
                }
            )

        return {
            "title": metadata.get("title", ""),
            "metadata": metadata,
            "sections": sections,
            "images": [
                {"path": f"images/{img['filename']}", "filename": img["filename"]}
                for img in images
            ],
            "stats": {
                "total_sections": len(sections),
                "headings": sum(1 for s in sections if s.get("type") == "heading"),
                "paragraphs": sum(1 for s in sections if s.get("type") == "paragraph"),
                "tables": sum(1 for s in sections if s.get("type") == "table"),
                "code_blocks": sum(1 for s in sections if s.get("type") == "code"),
                "math_blocks": sum(1 for s in sections if s.get("type") == "math"),
                "images": len(images),
            },
        }
